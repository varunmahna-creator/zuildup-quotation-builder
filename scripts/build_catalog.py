"""Build catalog_v2.json (no-tier) from the customer-facing quote sheet DOCX.

Schema per item:
  {
    "id":          "stable-snake_case-id",
    "category":    "structure" | "bathroom" | ... (canonical key)
    "category_label": "Structure" | "Bathroom & Toilet" | ...   # display label
    "label":       "Sanitary Ware & CP Fitting"
    "rate":        35000          # number, 0 if N/A (descriptive-only)
    "rate_text":   "₹35,000 per Bathroom"   # display string (with unit)
    "unit":        "per_bathroom" | "per_sqft" | ... | "fixed" | "cap" | "descriptive"
    "brands":      ["Jaquar", "Hindware Italian"]
    "scope":       ["full"] | ["full","structure_only"]
    "description": "Long DOCX paragraph"
    "notes":       optional caveats (e.g. "Pro-rata variation applies")
  }

The new (no-tier) catalog is strictly derived from Customer_Facing_Quote_Sheet.docx tables 5–18.
Brand typo fixes from Phase 1's bugs.md still apply (Jaguar→Jaquar, etc.).
"""
import json, re
from pathlib import Path
from docx import Document

QB = Path("/opt/ocplatform/workspace/zuildup/quotation-builder")
SRC = None  # read from stdin
OUT = QB / "catalog/catalog.json"

# ---- Category mapping ----
# DOCX section header -> (canonical_key, display_label, scope)
# Some DOCX tables share a category bucket. table_5/6/7 are all "Structure"-ish in the customer sense.
# We split a few:
#   - table_5  Design & Drawings -> design_drawings
#   - table_6  Structure          -> structure
#   - table_7  (concrete + others, structurally still Structure category in display)
#   - table_8  BATHROOM           -> bathroom
#   - table_9  Kitchen            -> kitchen
#   - table_10 Doors/Windows/Wardrobe
#   - table_11 Flooring
#   - table_12 Electrical Work
#   - table_13 Water Management
#   - table_14 Ceiling & Elevation
#   - table_15 Safety & Security
#   - table_16 Paint & Polish
#   - table_17 General Aspects
#   - table_18 (waterproofing/concrete/lift)  -> general (structurally a "general aspects" tail)

CATEGORY_FROM_TABLE = {
    5:  ("design_drawings", "Design & Drawings",       ["full"]),
    6:  ("structure",       "Structure",               ["full", "structure_only"]),
    7:  ("structure",       "Structure",               ["full", "structure_only"]),
    8:  ("bathroom",        "Bathroom & Toilet",       ["full"]),
    9:  ("kitchen",         "Kitchen",                 ["full"]),
    10: ("doors_windows",   "Doors, Windows & Wardrobe",["full"]),
    11: ("flooring",        "Flooring",                ["full"]),
    12: ("electrical",      "Electrical Work",         ["full"]),
    13: ("water",           "Water Management",        ["full", "structure_only"]),
    14: ("ceiling",         "Ceiling & Elevation",     ["full"]),
    15: ("safety",          "Safety & Security",       ["full"]),
    16: ("paint",           "Paint & Polish",          ["full"]),
    17: ("general",         "General Aspects",         ["full", "structure_only"]),
    18: ("general",         "General Aspects",         ["full", "structure_only"]),
}

# ---- Brand typo fixes (carry from Phase 1 bugs.md) ----
BRAND_FIXES = {
    "Jaguar": "Jaquar",
    "Wilroy & Bosch": "Villeroy & Boch",
}

# ---- Manual override map: id -> {brands, rate, unit, rate_text} ----
# Curated by reading the customer-facing DOCX directly. This is the source of truth
# for cases where the heuristic extractor mis-parses brands or rates.
OVERRIDES = {
    # bathroom — sanitary ware mentions Jaquar / Hindware Italian
    "bathroom.sanitary_ware_and_cp_fitting": {
        "brands": ["Jaquar", "Hindware (Italian Collection)"],
    },
    "bathroom.bathroom_geyser": { "brands": ["Crompton", "Bajaj"] },
    "bathroom.exhaust_fan":     { "brands": ["Havells"] },
    "kitchen.kitchen_wall_cladding_non_visible_areas": { "brands": [] },  # no brand, just "Commercial / Waste tiles"
    "kitchen.sink_and_faucet":  { "brands": [] },  # description-only
    "kitchen.kitchen_geyser":   { "brands": ["Crompton", "Bajaj"] },
    "kitchen.exhaust_fan":      { "brands": ["Havells"] },
    "kitchen.hob_and_chimney":  { "brands": ["Faber", "KAFF"] },
    "doors_windows.internal_and_bathroom_doors": { "brands": ["Sumai"] },
    "doors_windows.internal_and_bathroom_door_lock": { "brands": ["Godrej", "Ozone", "Yale"] },
    "doors_windows.sliding_door_windows":        { "brands": ["Prominence", "EITI"] },
    "doors_windows.modular_walkin_wardrobe":     { "brands": ["Godrej", "Ozone"] },
    "flooring.lift_fa_ade":                      { "brands": [] },
    "electrical.electrical_wire":                { "brands": ["Finolex", "Polycab"] },
    "electrical.electrical_wire_specification":  { "brands": [] },
    "electrical.conduit_pipes_boxes":            { "brands": ["Astral", "AKG"] },
    "electrical.downlighters":                   { "brands": ["Philips", "Osram"] },
    "electrical.cove_strip_lights":              { "brands": ["Philips", "Osram"] },
    "electrical.mcb_elcb":                       { "brands": [] },
    "structure.curing":                          { "brands": ["Sika Antisol"] },
    "structure.waterproofing":                   { "brands": ["Sika", "Dr. Fixit LW+"] },
    "structure.anti_termite_treatment":          { "brands": ["Bayer", "Terminator (Pidilite)"] },
    "water.underground_water_tank":              { "brands": ["Dr. Fixit", "Sika"] },
    "ceiling.false_ceiling":                     { "brands": ["Sakarni POP"] },
    "safety.cctv_camera":                        { "brands": ["CP Plus"], "rate": 50000, "unit": "cap", "rate_text": "₹50,000 (cap)" },
    "safety.video_door_phone":                   { "brands": ["CP Plus", "Alba"], "rate": 50000, "unit": "cap", "rate_text": "₹50,000 (cap)" },
    "paint.internal_wall_paint":                 { "brands": ["Asian Paints (Apcolite Premium Emulsion)"] },
    "paint.exterior_wall_paint":                 { "brands": ["Berger (Weathercoat Glow)"] },
    "paint.ceiling_paint":                       { "brands": ["Asian Paints (Apcolite Premium Emulsion)"] },
    "paint.putty":                               { "brands": ["JK Birla"] },
    "paint.polish":                              { "brands": [] },
    "general.main_gate":                         { "brands": [], "rate": 250000, "unit": "cap", "rate_text": "₹2,50,000 (cap, includes side gate)" },
    "general.lift_machine":                      { "brands": ["Schindler", "Kone"], "rate": 1000000, "unit": "cap", "rate_text": "₹10,00,000 (cap)" },
    # ---- P1.2: explicit rate_text for items the parse_rate regex can't catch ----
    "bathroom.shower_partition_cubicles":         { "rate": 10000, "unit": "per_bathroom", "rate_text": "upto ₹10,000 per bathroom" },
    "bathroom.bathroom_accessories":              { "brands": ["Jaquar"], "rate": 12000, "unit": "per_bathroom", "rate_text": "upto ₹12,000 per bathroom" },
    "bathroom.bathroom_flooring":                 { "rate": 40, "unit": "per_sqft", "rate_text": "Anti-skid tile cap ₹40 per sq.ft." },
    "bathroom.cpvc_fittings":                     { "brands": ["Jaquar", "Kohler"], "rate": 20000, "unit": "per_bathroom", "rate_text": "upto ₹20,000 per bathroom" },
    "kitchen.cpvc_fittings":                      { "brands": ["Jaquar", "Kohler"], "rate": 20000, "unit": "per_bathroom", "rate_text": "upto ₹20,000 per bathroom" },
    "kitchen.modular_kitchen":                    { "brands": ["Ozone"], "rate": 250000, "unit": "cap", "rate_text": "upto ₹2,50,000 per kitchen" },
    "doors_windows.main_entry_door":              { "rate": 20000, "unit": "per_door", "rate_text": "upto ₹20,000 per door (Teak/Mikasa/Sumai)" },
    "doors_windows.main_door_lock":               { "brands": ["Godrej", "Ozone", "Yale"], "rate": 12000, "unit": "cap", "rate_text": "upto ₹12,000 (Godrej/Ozone/Yale)" },
    "doors_windows.terrace_door":                 { "rate": 26000, "unit": "fixed", "rate_text": "₹26,000 (Tata Parvesh GS)" },
    "flooring.floor_flooring":                    { "rate": 250, "unit": "per_sqft", "rate_text": "Italian Marble — ₹250 per sq.ft." },
    "flooring.balcony_flooring":                  { "rate": 100, "unit": "per_sqft", "rate_text": "Granite Stone — ₹100 per sq.ft." },
    "flooring.terrace_flooring":                  { "rate": 40, "unit": "per_sqft", "rate_text": "Anti-skid tile — ₹40 per sq.ft." },
    "flooring.lift_fa_ade":                       { "rate": 100, "unit": "per_sqft", "rate_text": "Tiles/Stone/Wooden Look — ₹100 per sq.ft." },
    "electrical.switch_sockets":                  { "brands": ["Havells", "Anchor", "LeGrand"], "rate": 50000, "unit": "per_floor", "rate_text": "₹50,000 per floor (LeGrand white & dark grey)" },
    "electrical.ceiling_fans":                    { "brands": ["Havells"], "rate": 1800, "unit": "per_fan", "rate_text": "₹1,800 per fan (Havells)" },
    "electrical.pillar_fancy_light":              { "rate": 2500, "unit": "per_piece", "rate_text": "₹2,500 per light" },
    "water.overhead_water_tank":                  { "brands": ["Astral"], "rate": 8500, "unit": "cap", "rate_text": "1×1000 L Triple-Layer per dwelling, upto ₹8,500" },
    "water.water_motor":                          { "brands": ["Crompton Greaves"], "rate": 8500, "unit": "per_floor", "rate_text": "₹8,500 per floor (Crompton Greaves)" },
    "general.staircase_balcony_railing":          { "rate": 400, "unit": "per_sqft", "rate_text": "MS Steel designer railing — ₹400 per sq.ft." },
    "structure.steel":                            { "brands": ["Rathi Steel 500FE"], "rate": 55000, "unit": "per_sqft", "rate_text": "@5kg/sqft, ₹55,000/MT cap (Rathi 500FE)" },
    "structure.cement":                           { "brands": ["Ultratech", "ACC"], "rate": 380, "unit": "per_piece", "rate_text": "upto ₹380 per bag (Ultratech/ACC)" },
    "structure.bricks":                           { "rate": 7, "unit": "per_piece", "rate_text": "upto ₹7.50 per brick (A Class)" },
}

# ---- Helpers ----

def slugify(s):
    s = re.sub(r'[^A-Za-z0-9]+', '_', s).strip('_').lower()
    s = re.sub(r'_+', '_', s)
    return s

def parse_rate(spec_text, label):
    """Heuristic: pull the first explicit money amount out of the text.
    Returns (rate_number, rate_text, unit).

    Strategy:
      1. Look for "X Lacs" / "X Lakhs" / "X Lac" / "X Crore" / "X Cr" → multiply.
      2. Look for "₹/Rs/INR <num>" or bare "<num>/- per <unit>" or "<num> per <unit>".
      3. Fallback: descriptive (no rate).
    """
    blob = (spec_text or "") + " " + (label or "")
    rate = 0

    # 1. Lacs / Lakh / Crore form (e.g. "INR 10 Lacs", "Rs 25 Lakh", "1.5 Crore")
    lacs = re.search(r'(?:₹|Rs\.?|INR)?\s*([\d.]+)\s*(?:lac|lakh|lakhs|lacs)\b', blob, re.I)
    cr = re.search(r'(?:₹|Rs\.?|INR)?\s*([\d.]+)\s*(?:crore|cr)\b', blob, re.I)
    if lacs:
        try: rate = int(float(lacs.group(1)) * 100000)
        except: rate = 0
    if cr and rate == 0:
        try: rate = int(float(cr.group(1)) * 10000000)
        except: rate = 0

    # 2. Money with explicit currency marker. Indian thousands: 35,000 / 2,50,000.
    if rate == 0:
        m = re.search(r'(?:₹|Rs\.?|INR)\s*([\d,]+)(?:/-)?', blob)
        if m:
            try: rate = int(m.group(1).replace(',', ''))
            except: rate = 0

    # 3. Bare number followed by "/- per X" or "per X" without currency
    if rate == 0:
        m = re.search(r'\b(\d{2,3}(?:,\d{2,3})+|\d{2,5})/-?\s*per\b', blob, re.I)
        if not m:
            m = re.search(r'\b(\d{2,3}(?:,\d{2,3})+|\d{2,5})\s*per\s*(?:bathroom|sq\.?\s*ft|floor|piller|pillar|piece|fan|door)\b', blob, re.I)
        if m:
            try: rate = int(m.group(1).replace(',', ''))
            except: rate = 0

    # 4. Compound rate+cap: Steel pattern "@ 5Kg per Sq Ft … Rs. 55000/MT"
    if rate == 0:
        m = re.search(r'@\s*(\d+(?:\.\d+)?)\s*Kg\s*per\s*Sq\.?\s*Ft.*?Rs\.?\s*(\d+(?:,\d+)*)\s*/\s*MT', blob, re.I)
        if m:
            try:
                kg = m.group(1)
                cap = int(m.group(2).replace(',', ''))
                rate = cap
                # store special rate_text via override path (set unit; rate_text built below)
                # use sentinel: store in _compound for rate_text section to detect
                _compound = (kg, cap)
            except Exception:
                _compound = None
        else:
            _compound = None
    else:
        _compound = None

    # 5. Decimal rate: "Rs. 7.50 per brick"
    if rate == 0:
        m = re.search(r'Rs\.?\s*(\d+\.\d+)\s*per\s*(\w+)', blob, re.I)
        if m:
            try:
                rate = int(round(float(m.group(1))))  # store rounded int (catalog rate is int)
                _decimal = float(m.group(1))
                _decimal_unit = m.group(2)
            except Exception:
                _decimal = None
                _decimal_unit = None
        else:
            _decimal = None
            _decimal_unit = None
    else:
        _decimal = None
        _decimal_unit = None

    # 6. No-per form: "Rs. 250/- sq ft" (suffix without explicit per)
    if rate == 0:
        m = re.search(r'Rs\.?\s*(\d+(?:,\d+)*)\s*/-?\s*sq\.?\s*ft', blob, re.I)
        if m:
            try:
                rate = int(m.group(1).replace(',', ''))
                _noper_sqft = True
            except Exception:
                _noper_sqft = False
        else:
            _noper_sqft = False
    else:
        _noper_sqft = False

    # Unit detection
    txt_low = blob.lower()
    if 'per bathroom' in txt_low: unit = 'per_bathroom'
    elif 'per sq' in txt_low or 'per sq ft' in txt_low or 'per sqft' in txt_low: unit = 'per_sqft'
    elif 'per floor' in txt_low: unit = 'per_floor'
    elif 'per piller' in txt_low or 'per pillar' in txt_low: unit = 'per_pillar'
    elif 'per piece' in txt_low: unit = 'per_piece'
    elif 'per fan' in txt_low: unit = 'per_fan'
    elif 'per door' in txt_low: unit = 'per_door'
    elif 'per kg' in txt_low or 'kg per sqft' in txt_low: unit = 'per_sqft'  # steel
    elif rate > 0 and ('cap' in txt_low or 'upto' in txt_low or 'up to' in txt_low or 'within' in txt_low or 'budget' in txt_low or 'specified value' in txt_low or 'includes' in txt_low):
        unit = 'cap'
    elif rate > 0:
        unit = 'fixed'
    else:
        unit = 'descriptive'

    # rate_text: pretty render
    if rate > 0:
        # Indian format
        s = str(rate)
        if len(s) > 3:
            l = s[-3:]; r = s[:-3]
            parts = []
            while len(r) > 2:
                parts.insert(0, r[-2:]); r = r[:-2]
            if r: parts.insert(0, r)
            rate_str = "₹" + ",".join(parts) + "," + l
        else:
            rate_str = "₹" + s

        unit_label = {
            'per_bathroom': 'per bathroom', 'per_sqft': 'per sq.ft.', 'per_floor': 'per floor',
            'per_pillar': 'per pillar', 'per_piece': 'per piece', 'per_fan': 'per fan',
            'per_door': 'per door', 'cap': '(cap)', 'fixed': '',
        }.get(unit, '')
        rate_text = (rate_str + " " + unit_label).strip()
        # Override for compound/decimal/no-per special cases
        if _compound is not None:
            kg, cap = _compound
            # pretty-print cap with Indian commas
            cs = str(cap); l3 = cs[-3:]; rr = cs[:-3]; pp = []
            while len(rr) > 2:
                pp.insert(0, rr[-2:]); rr = rr[:-2]
            if rr: pp.insert(0, rr)
            cap_str = "₹" + (",".join(pp) + "," + l3 if pp else l3)
            rate_text = f"@{kg}kg/sqft, {cap_str}/MT cap"
            unit = "per_sqft"
        elif _decimal is not None and _decimal_unit:
            rate_text = f"₹{_decimal:.2f} per {_decimal_unit.lower()}"
            unit = "per_piece" if _decimal_unit.lower() in ("brick", "piece", "fan", "door", "light") else unit
        elif _noper_sqft:
            rate_text = f"{rate_str} per sq.ft."
            unit = "per_sqft"
    else:
        rate_text = ""

    return rate, rate_text, unit


def extract_brands(spec_text):
    """Pull brand names from the first line (or sometimes embedded in description).
    A 'brand' for our purposes is a TitleCase or all-caps token like 'Jaquar', 'Asian Paints',
    'CP Plus', 'Dr. Fixit'. NOT money strings, NOT generic words.
    """
    if not spec_text: return []
    # Look only at first line — that's where ZuildUp puts brands in this DOCX
    first_line = spec_text.split('\n')[0].strip()
    # Strip trailing parenthetical comments like "(CP Plus or equivalent)" — actually those CONTAIN brands, keep.
    if not first_line: return []
    # Reject if first line is a money string or pure description
    if re.match(r'^(₹|Rs\.?|INR|\d)', first_line) and 'per' in first_line.lower():
        # rate line, not brand line
        return []
    # If the line is too long it's probably a description, not a brand line — but it can still mention brands inside parens
    work = first_line
    # Look for an explicit "(BrandA / BrandB or equivalent)" parenthetical first
    paren = re.search(r'\(([^)]*\b(?:make|or equivalent|by|brand)[^)]*)\)', first_line, re.I)
    if paren:
        work = paren.group(1)

    cands = []
    # Split on / , or  "or equivalent" or "and" — common separators
    parts = re.split(r'\s*[/,]\s*|\s+or\s+equivalent\b|\s+(?:and|by|make|equivalent)\s+', work, flags=re.I)
    for p in parts:
        p = p.strip().rstrip('.,;:').strip()
        if not p: continue
        if len(p) < 2 or len(p) > 40: continue
        # must start with an uppercase letter (or "Dr.")
        if not re.match(r'^(?:Dr\.\s*)?[A-Z]', p): continue
        # reject if contains digits (rate strings) or currency
        if re.search(r'\d|₹|Rs\.?|INR', p): continue
        # reject if it has obvious sentence words
        if re.search(r'\b(per|shall|will|the|a|an|of|with|for|to|from|using|including|provided|made|finish|design|grade|copper|sand|aggregate|ratio)\b', p, re.I):
            continue
        # reject if it looks like a generic noun phrase (lowercase predominance)
        words = p.split()
        if len(words) > 5: continue
        cap_ratio = sum(1 for w in words if w and w[0].isupper()) / max(1, len(words))
        if cap_ratio < 0.5: continue
        cands.append(p)

    # apply fixes + dedupe
    seen = []
    for b in cands:
        b = BRAND_FIXES.get(b, b)
        if b and b not in seen:
            seen.append(b)
    return seen


def fix_label(label):
    """Capitalisation + typo fixes."""
    fixes = {
        'SInk and Faucet': 'Sink and Faucet',
        'BATHROOM': 'Bathroom',
    }
    return fixes.get(label, label)


def main():
    # FS cache wedged on this shell's mount namespace. open() / stat() / md5sum can't see
    # the DOCX, but cat-via-openat() can. Read via stdin (piped-in cat).
    import io, sys
    buf = io.BytesIO(sys.stdin.buffer.read())
    print(f"read {buf.getbuffer().nbytes} bytes from stdin", file=sys.stderr)
    doc = Document(buf)
    items = []

    for ti, (cat_key, cat_label, scope) in CATEGORY_FROM_TABLE.items():
        if ti >= len(doc.tables): continue
        tbl = doc.tables[ti]
        rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
        # row 0 = category header, skip
        for ri in range(1, len(rows)):
            cells = rows[ri]
            if len(cells) < 2: continue
            label_cell = cells[0]
            spec_cell  = cells[1]
            if not label_cell or label_cell == spec_cell:
                continue  # likely sub-header
            # First line of spec_cell is often the rate / brand line; rest is description
            lines = [l.strip() for l in spec_cell.split('\n') if l.strip()]
            first = lines[0] if lines else ""
            desc  = "\n".join(lines[1:]) if len(lines) > 1 else ""
            # If first line is itself the description (no rate, no brand pattern), keep all in description
            if not desc:
                desc = first
                first = ""

            label = fix_label(re.sub(r'\s+', ' ', label_cell))
            rate, rate_text, unit = parse_rate(first + " " + desc, label)
            brands = extract_brands(first if first else spec_cell)

            item_id = f"{cat_key}.{slugify(label)}"
            item = {
                "id": item_id,
                "category": cat_key,
                "category_label": cat_label,
                "label": label,
                "rate": rate,
                "rate_text": rate_text,
                "unit": unit,
                "brands": brands,
                "scope": scope,
                "description": ((first + ("\n" + desc if desc else "")).strip() if first else desc.strip()),
            }
            # apply manual overrides (curated from DOCX read-through)
            ov = OVERRIDES.get(item_id, {})
            for k, v in ov.items():
                item[k] = v
            # dedupe — same id within table run
            if not any(i['id'] == item['id'] for i in items):
                items.append(item)

    catalog = {
        "_meta": {
            "version": "2.0.0",
            "built": "2026-04-28",
            "source": "src_docx/Customer_Facing_Quote_Sheet.docx (Customer: Ms. Rajkumari Kamboj reference quote)",
            "schema": "no-tier; single rate + single brands per item; team enters cost-per-sqft per quote",
            "tier_system": "removed in Phase 2 P0.1",
            "items_total": len(items),
        },
        "items": items,
    }
    # FS wedge on workspace path — write to /tmp then we'll cat it into place via tee.
    tmp_out = "/tmp/_catalog_v2.json"
    with open(tmp_out, "w") as f:
        f.write(json.dumps(catalog, indent=2, ensure_ascii=False))
    print(f"WROTE {tmp_out}")
    print(f"  items: {len(items)}")
    print(f"  size : {OUT.stat().st_size} bytes")
    print(f"\nBy category:")
    from collections import Counter
    by = Counter(i['category'] for i in items)
    for c, n in by.most_common():
        print(f"  {c}: {n}")
    print(f"\nUnit distribution:")
    by_u = Counter(i['unit'] for i in items)
    for u, n in by_u.most_common():
        print(f"  {u}: {n}")
    print(f"\nItems with explicit rate (>0): {sum(1 for i in items if i['rate'] > 0)}")
    print(f"Items with brands: {sum(1 for i in items if i['brands'])}")
    print(f"\nSample items:")
    for i in items[:5]:
        print(f"  - {i['id']:45s}  rate={i['rate']:>8}  unit={i['unit']:<14} brands={i['brands']}")
    print(f"  ... and {len(items)-5} more")

if __name__ == '__main__':
    main()
